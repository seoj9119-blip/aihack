import io
from dataclasses import dataclass
from pathlib import Path

import docx
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
from xml.sax.saxutils import escape

FONT_PATH = Path(__file__).resolve().parent.parent / "assets" / "fonts" / "NanumGothic-Regular.ttf"
FONT_NAME = "NanumGothic"

if FONT_NAME not in pdfmetrics.getRegisteredFontNames():
    pdfmetrics.registerFont(TTFont(FONT_NAME, str(FONT_PATH)))


@dataclass
class Block:
    kind: str  # "h1" | "h2" | "quote" | "bullet" | "p"
    text: str


def _parse_markdown(content: str) -> list[Block]:
    blocks: list[Block] = []
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            blocks.append(Block("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(Block("h2", line[3:].strip()))
        elif line.startswith("> "):
            blocks.append(Block("quote", line[2:].strip()))
        elif line.startswith("- ") or line.startswith("* "):
            blocks.append(Block("bullet", line[2:].strip()))
        else:
            blocks.append(Block("p", line))
    return blocks


def markdown_to_pdf(content: str) -> bytes:
    blocks = _parse_markdown(content)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    style_h1 = ParagraphStyle("h1", fontName=FONT_NAME, fontSize=20, leading=26, spaceAfter=12)
    style_h2 = ParagraphStyle("h2", fontName=FONT_NAME, fontSize=15, leading=20, spaceAfter=8, spaceBefore=6)
    style_quote = ParagraphStyle(
        "quote", fontName=FONT_NAME, fontSize=10, leading=15, textColor="#555555", spaceAfter=8
    )
    style_p = ParagraphStyle("p", fontName=FONT_NAME, fontSize=11, leading=16, spaceAfter=6)
    style_bullet = ParagraphStyle("bullet", fontName=FONT_NAME, fontSize=11, leading=16)

    story = []
    bullet_buffer: list[str] = []

    def flush_bullets():
        if bullet_buffer:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(escape(text), style_bullet)) for text in bullet_buffer],
                    bulletType="bullet",
                )
            )
            story.append(Spacer(1, 6))
            bullet_buffer.clear()

    for block in blocks:
        if block.kind == "bullet":
            bullet_buffer.append(block.text)
            continue
        flush_bullets()
        if block.kind == "h1":
            story.append(Paragraph(escape(block.text), style_h1))
        elif block.kind == "h2":
            story.append(Paragraph(escape(block.text), style_h2))
        elif block.kind == "quote":
            story.append(Paragraph(escape(block.text), style_quote))
        else:
            story.append(Paragraph(escape(block.text), style_p))
    flush_bullets()

    doc.build(story)
    return buffer.getvalue()


def markdown_to_docx(content: str) -> bytes:
    blocks = _parse_markdown(content)
    document = docx.Document()

    for block in blocks:
        if block.kind == "h1":
            document.add_heading(block.text, level=1)
        elif block.kind == "h2":
            document.add_heading(block.text, level=2)
        elif block.kind == "quote":
            p = document.add_paragraph(block.text)
            p.style = document.styles["Intense Quote"]
        elif block.kind == "bullet":
            document.add_paragraph(block.text, style="List Bullet")
        else:
            document.add_paragraph(block.text)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
